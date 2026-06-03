#!/usr/bin/env python3
"""Generate Word walkthrough script for ADGM Command Centre demo."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "docs" / "ADGM-Command-Centre-Walkthrough-Script.docx"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_title(doc: Document, text: str, level: int = 0) -> None:
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)


def add_meta(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_label(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(label.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x77)


def add_say(doc: Document, text: str) -> None:
    add_label(doc, "Say")
    for block in text.strip().split("\n\n"):
        p = doc.add_paragraph(block.strip())
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.italic = True


def add_do(doc: Document, text: str) -> None:
    add_label(doc, "Do / Show")
    p = doc.add_paragraph(text.strip())
    p.paragraph_format.left_indent = Inches(0.25)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_transition(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("→ Transition: ")
    run.bold = True
    p.add_run(text.strip())


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    set_doc_defaults(doc)

    # Title page
    add_title(doc, "Personal AI for Rajiv Sehgal")
    add_title(doc, "Command Centre — Product Walkthrough Script", level=1)
    doc.add_paragraph()
    add_meta(
        doc,
        [
            "Audience: ADGM leadership, IT, strategy stakeholders",
            "Duration: 12–15 minutes (full) · 5–6 minutes (short cut)",
            "Demo persona: Rajiv Sehgal, Chief Strategy Officer, ADGM",
            "Product: Personal AI Command Centre (web application)",
        ],
    )
    doc.add_page_break()

    # Before you start
    add_title(doc, "Before You Start", level=1)
    doc.add_heading("Environment checklist", level=2)
    for item in [
        "Browser at full width (or iPhone simulator for mobile proof).",
        "Language: EN first; switch to ع (Arabic) once near the end.",
        "Logged in as executive view; demo data loaded.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_label(doc, "Opening line (30 sec)")
    add_say(
        doc,
        """Good [morning/afternoon]. Today I'll walk you through Personal AI for Rajiv Sehgal — not a generic chatbot, but a command centre built for how a Chief Strategy Officer actually works: nine departments, global markets, regulation, institutional knowledge, and executive briefings — in one place, in English and Arabic.""",
    )

    add_label(doc, "Set the problem (30 sec)")
    add_say(
        doc,
        """CSOs drown in fragmented tools — email, SharePoint, Bloomberg, CRM, board packs. Decisions slow down because context is scattered. This product pulls signals, performance, market, regulation, and documents into one surface, with a personal AI agent that orchestrates five specialists behind a single question.""",
    )

    # Screen overview table
    doc.add_page_break()
    add_title(doc, "Screen Overview", level=1)
    add_table(
        doc,
        ["#", "Screen", "Route", "Purpose"],
        [
            ["1", "Executive Home", "/dashboard", "Morning command view — priorities, signals, AI insight"],
            ["2", "Ask Personal AI", "/chat", "Natural-language Q&A across all intelligence"],
            ["3", "Performance", "/performance", "Live health of 9 departments + escalations"],
            ["4", "Market Intelligence", "/market", "Capital flows, benchmarks, opportunities"],
            ["5", "Regulatory", "/regulatory", "Global regulatory change monitor"],
            ["6", "Knowledge Base", "/knowledge", "Approved documents — search, upload, preview"],
            ["7", "Briefings", "/briefings", "Auto-generate executive brief formats"],
            ["8", "Architecture", "/architecture", "Agents, orchestration, integrations, delivery"],
            ["9", "Settings", "/settings", "Profile, language, theme, AI behaviour"],
        ],
    )

    acts = [
        {
            "act": "Act 1 — The Command Centre (Executive Home)",
            "route": "/dashboard",
            "time": "~2 minutes",
            "purpose": "Single landing page for the CSO: overnight summary, live market pulse, top risks, and shortcuts into AI and Performance.",
            "features": [
                ("Market ticker", "Scrolls live-style KPIs so you see market context without leaving home."),
                ("Greeting hero", "Abu Dhabi time/date, welcome message, quick actions (morning briefing, Performance)."),
                ("Priority signals", "Most important issues today — metric, trend, summary; tap to drill in."),
                ("AI insight card", "Cross-department patterns (e.g. IT risk linked to Sales deal risk)."),
                ("Ask chips", "Pre-written questions that open Chat with one click."),
                ("Disclaimer", "Clarifies demo content vs official ADGM policy."),
            ],
            "do": "Land on dashboard. Let the market ticker scroll. Scroll through 2–3 signal cards. Point to AI insight card. Scroll to Ask chips (do not click yet).",
            "say": """This is Executive Home — Rajiv's morning command deck on Al Maryah Island, Abu Dhabi time.

The ticker gives live market context without opening another terminal.

The hero is personalised: good morning Rajiv, what happened overnight across strategic priorities and nine teams. Two primary actions: open a morning briefing in chat, or jump straight to Performance.

Below that: Today's priority signals — not a news feed, but curated executive signals. Each card shows a label, headline, metric, trend line, and short narrative. Green, amber, or red tone tells you urgency at a glance. I can tap any card to drill into Performance or the relevant area.

Notice each signal is decision-oriented — 'what do I need to know?' not 'what happened in the market generically?'

This is pattern detection across departments. Here, an IT integration slip is tied to a stalled sovereign-fund deal in Sales — something you might miss if you only read one dashboard. Investigate sends that straight into the AI with the right context.

Suggested questions at the bottom are one-click entry points into the assistant. The disclaimer is clear: illustrative executive scenarios; official ADGM documents always prevail.""",
            "transition": "Home answers: What matters this morning? Next: How is the organisation performing right now?",
        },
        {
            "act": "Act 2 — Performance (Nine departments, one view)",
            "route": "/performance",
            "time": "~2.5 minutes (short: 60 sec — escalations + one dept)",
            "purpose": "Organisation-wide performance cockpit: escalations, org momentum, all departments with RAG status.",
            "features": [
                ("Escalation panel", "Urgent items across HR, Legal, Procurement, Sales — tap to open department."),
                ("Org overview", "Donut by health; momentum chart; D33, hours saved, briefing coverage."),
                ("Department cards", "Summary, RAG, risk count, sparkline; click for detail."),
                ("Department detail", "KPIs, achievements, blockers, risks, recommended actions."),
            ],
            "do": "Pause on Escalation panel. Click one escalation or point. Scroll to donut/momentum. Open one department (Sales or HR), walk detail, then Back.",
            "say": """Performance is the CSO's operating system for nine departments — HR, Legal, Procurement, Sales, IT, and others — all on one screen.

Escalations are what needs leadership today: attrition thresholds, contract expiry, regulatory filings, stalled mandates. I tap one — for example the AED 90M sovereign-fund mandate — and we go into that department's detail.

Below escalations: organisational health. The donut shows green, amber, red departments. The momentum index shows trajectory, plus D33 alignment, hours saved per week, and briefing coverage.

Each department card shows summary, RAG status, high-risk count, and sparkline. Rajiv manages by exception — scan reds and ambers.

Department detail: KPIs, achievements, blockers, risks, and numbered recommended actions. Performance answers: Who needs me today, and are we winning?""",
            "transition": "Performance is internal. Strategy also needs the external picture — capital, competitors, and D33 opportunities.",
        },
        {
            "act": "Act 3 — Market Intelligence",
            "route": "/market",
            "time": "~2 minutes (short: briefing + one chart)",
            "purpose": "Strategic market view for Abu Dhabi / ADGM: brief, competitive position, capital flows, D33 opportunities.",
            "features": [
                ("Morning briefing", "Overnight bullets: GCC flows, DIFC, AI/VC, yields."),
                ("Radar chart", "ADGM vs selected global financial centre."),
                ("Capital flow viz", "Net inflows by region toward Abu Dhabi."),
                ("12-dimension benchmark", "ADGM vs global centres per dimension."),
                ("Investment opportunities", "Ranked by D33 fit score."),
                ("Differentiation", "Actionable ideas for Abu Dhabi to lead."),
            ],
            "do": "Morning briefing card → radar + capital flow → benchmark + opportunities → differentiation.",
            "say": """Market Intelligence answers: where is capital going, and where can Abu Dhabi win?

The morning briefing is auto-generated at 06:00 GST — GCC flows, competitor moves, AI inflows, yield moves.

Radar and capital flow compare ADGM versus global centres. The benchmark scores ADGM on twelve dimensions. Investment opportunities are ranked by D33 fit — AI infrastructure, private credit, tokenised assets, transition finance.

The bottom section is strategic differentiation — actionable ideas, not analyst fluff.""",
            "transition": "Markets move fast; regulation moves faster. That's a separate dedicated monitor.",
        },
        {
            "act": "Act 4 — Regulatory",
            "route": "/regulatory",
            "time": "~1 minute (short: 30 sec)",
            "purpose": "Regulatory change monitor — tracks major bodies, scores relevance, links to AI impact analysis.",
            "features": [
                ("Jurisdiction summary", "12 bodies monitored; Policy AI monitoring."),
                ("All / High filter", "Every alert or high-relevance only."),
                ("Regulatory cards", "Body, region, topic, relevance, impact summary."),
                ("Tap → Chat", "Pre-filled 'analyse impact on ADGM' prompt."),
            ],
            "do": "Summary card, chips, All/High filter. Tap one card to chat.",
            "say": """Regulatory tracks twelve jurisdictions — FSRA, MAS, FCA, SEC, and more — ingested within hours.

Filter High relevance when time is short. Each card explains why it matters to Abu Dhabi. Tap one to get full AI impact analysis in chat.""",
            "transition": "Intelligence is only trustworthy if it's grounded in approved institutional knowledge.",
        },
        {
            "act": "Act 5 — Knowledge Base",
            "route": "/knowledge",
            "time": "~1.5 minutes (short: search + mention upload)",
            "purpose": "Institutional document hub — search, filter, upload, preview; powers grounded AI answers.",
            "features": [
                ("Search bar", "Find documents; can route to chat."),
                ("Try chips", "One-click common CSO queries."),
                ("Add document", "File picker → side panel → upload progress."),
                ("Category filters", "Narrow list by document type."),
                ("Document table", "Title, topic, date, status, actions."),
                ("Preview sheet", "Metadata and content summary."),
            ],
            "do": "Search, Try chips, filters, table. Optional: Add document flow.",
            "say": """Knowledge Base is approved institutional memory.

Search or use Try chips. Add document: pick a file, set category and date, see upload progress. Filter by category. Open any row to preview. This corpus is what the AI cites — not the open internet.""",
            "transition": "When Rajiv wants a briefing in seconds, not a document search — that's Briefings.",
        },
        {
            "act": "Act 6 — Briefings",
            "route": "/briefings",
            "time": "~1 minute (short: pre-meeting only)",
            "purpose": "Executive briefing generator — pick format, get structured narrative in seconds.",
            "features": [
                ("Pre-meeting brief", "Calendar context, under 30 seconds."),
                ("Board pack summary", "Decisions, not pages."),
                ("Stakeholder profile", "CRM + enrichment."),
                ("Policy impact analysis", "Regulatory change."),
                ("Strategic opportunity brief", "D33-scored."),
                ("Ministerial note", "Bilingual Arabic and English."),
            ],
            "do": "Show six format cards. Select Pre-meeting → Generate → skim output.",
            "say": """Briefing generator produces decision-ready output in seconds. Six formats from pre-meeting to ministerial notes.

Select Pre-meeting brief, hit Generate — structured narrative with headings and bullets, ready to refine in chat or paste into email. This is AI that ships executive work products.""",
            "transition": "Now where it all comes together — the Personal AI Agent.",
        },
        {
            "act": "Act 7 — Ask Personal AI (Chat)",
            "route": "/chat",
            "time": "~2 minutes (short: one suggestion)",
            "purpose": "Conversational front door; five agents orchestrated via LangGraph; one unified answer.",
            "features": [
                ("Chief of Staff AI", "Calendar, actions, board packs."),
                ("Strategy AI", "Market, benchmarking, D33."),
                ("Policy AI", "Regulation and impact."),
                ("Relationship AI", "Stakeholders and CRM."),
                ("Communications AI", "Speeches, bilingual notes."),
            ],
            "do": "Empty state → click suggestion or type: Brief me on my 3pm meeting tomorrow. Read opening of answer.",
            "say": """Ask Personal AI is conversational — plain executive language.

Behind one answer, five agents run in parallel: Chief of Staff, Strategy, Policy, Relationship, Communications. The user sees one unified brief.

Click a suggestion or type a question. Answers are structured for scanning. In production: Claude on local server; demo mode may show canned content. Sources tie to knowledge base.""",
            "transition": "For technical stakeholders — how it's built.",
        },
        {
            "act": "Act 8 — Architecture",
            "route": "/architecture",
            "time": "~1 minute (short: agents only or skip)",
            "purpose": "Technical view: agents, orchestration, integrations, eight-week delivery on Azure UAE North.",
            "features": [
                ("Five agent cards", "Role, capabilities, integrations."),
                ("LangGraph demo", "One query → multiple agents → one brief."),
                ("Integrations grid", "Bloomberg, CRM, SharePoint — Phase 1."),
                ("8-week plan", "Demo to go-live week 8."),
            ],
            "do": "Agent cards → orchestration example → integrations → timeline.",
            "say": """Architecture is transparent. Five agents with clear roles. LangGraph example: one delegation question fires four agents; one answer in under thirty seconds. Integrations and eight-week delivery plan. Hosted on Azure UAE North.""",
            "transition": "Finally — how Rajiv makes it his own.",
        },
        {
            "act": "Act 9 — Settings & Arabic",
            "route": "/settings",
            "time": "~45 seconds",
            "purpose": "Profile, agent routing, language (EN/AR + RTL), theme, response style.",
            "features": [
                ("Profile", "Name, email, title."),
                ("Smart routing", "Auto or manual agent selection."),
                ("Language EN / ع", "Full UI RTL in Arabic."),
                ("Response style", "Concise, detailed, executive."),
            ],
            "do": "Profile → smart routing → toggle ع in top bar → pause on one Arabic screen → back to EN.",
            "say": """Settings: profile, smart agent routing, theme, response style. Switch to Arabic — entire command centre flips RTL. Same product, bilingual executive experience.""",
            "transition": "",
        },
    ]

    for a in acts:
        doc.add_page_break()
        add_title(doc, a["act"], level=1)
        p = doc.add_paragraph()
        p.add_run("Route: ").bold = True
        p.add_run(a["route"])
        p.add_run("  |  Time: ").bold = True
        p.add_run(a["time"])

        doc.add_heading("What this screen does", level=2)
        doc.add_paragraph(a["purpose"])

        doc.add_heading("Features — what each does", level=2)
        ft = doc.add_table(rows=1 + len(a["features"]), cols=2)
        ft.style = "Table Grid"
        ft.rows[0].cells[0].text = "Feature"
        ft.rows[0].cells[1].text = "What it does"
        for i, (name, desc) in enumerate(a["features"], 1):
            ft.rows[i].cells[0].text = name
            ft.rows[i].cells[1].text = desc
        doc.add_paragraph()

        add_do(doc, a["do"])
        add_say(doc, a["say"])
        if a["transition"]:
            add_transition(doc, a["transition"])

    # Closing
    doc.add_page_break()
    add_title(doc, "Closing (60–90 sec)", level=1)
    add_say(
        doc,
        """To recap: Executive Home for the morning story. Performance for nine departments and escalations. Market and Regulatory for external strategy. Knowledge Base for grounded truth. Briefings for shipped work products. Chat for anything in natural language — five agents, one answer. Architecture for how we deliver in eight weeks on UAE infrastructure.

This is Personal AI for Rajiv Sehgal — strategic intelligence, market opportunity, policy insight, stakeholder readiness, and performance visibility — in one command centre.

I'm happy to go deeper on security, data residency, agent design, or a live scenario you choose.""",
    )

    doc.add_heading("Q&A — short answers", level=2)
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ["Is this live data?", "Demo uses illustrative data; architecture shows Bloomberg, Refinitiv, regulatory feeds in Phase 1."],
            ["Where is data hosted?", "Azure UAE North; SSO and RBAC in delivery plan."],
            ["Hallucinations?", "Answers grounded on approved knowledge base; official ADGM docs prevail."],
            ["Arabic quality?", "Full RTL UI; Communications AI for bilingual notes; Arabic QA in week 6."],
            ["Different from Copilot?", "Purpose-built CSO command centre + five domain agents + performance/regulatory/market modules."],
        ],
    )

    doc.add_page_break()
    add_title(doc, "Short Demo Cut (5–6 min)", level=1)
    for step in [
        "Home (45s) — ticker, one signal, AI insight",
        "Performance (60s) — escalations, one dept",
        "Market (45s) — morning briefing, one chart",
        "Briefings (45s) — generate pre-meeting",
        "Chat (90s) — one question",
        "Settings (20s) — Arabic toggle",
        "Close (30s) — recap",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Presenter cheat sheet", level=2)
    cheats = [
        ("HOME", "What matters this morning?"),
        ("PERFORMANCE", "Who needs me? Are we green?"),
        ("MARKET", "Where is capital? Where do we lead?"),
        ("REGULATORY", "What changed? What's the impact?"),
        ("KNOWLEDGE", "What's approved and searchable?"),
        ("BRIEFINGS", "Ship a brief in seconds."),
        ("CHAT", "One question, five agents, one answer."),
        ("ARCHITECTURE", "How we build and integrate."),
        ("SETTINGS", "English, Arabic, executive tone."),
    ]
    for screen, line in cheats:
        p = doc.add_paragraph()
        p.add_run(f"{screen}: ").bold = True
        p.add_run(line)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
